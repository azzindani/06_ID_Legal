"""
DOCX Extractor - Extract text from Word documents

File: document_parser/extractors/docx_extractor.py
"""

from typing import Dict, Any
from .base import BaseExtractor
from utils.logger_utils import get_logger


class DOCXExtractor(BaseExtractor):
    """Extract text from Word documents (.docx, .doc)"""
    
    SUPPORTED_EXTENSIONS = ['.docx', '.doc']
    EXTRACTOR_NAME = "docx"
    
    def __init__(self):
        super().__init__()
        self.logger = get_logger("DOCXExtractor")
    
    def _check_dependencies(self) -> bool:
        """Check if python-docx is available"""
        try:
            import docx
            return True
        except ImportError:
            return False
    
    def extract(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from Word document.
        
        Extracts paragraphs and tables.
        """
        self.validate_file(file_path)
        
        if not self.is_available():
            from ..exceptions import ExtractionError
            raise ExtractionError(
                file_path,
                "python-docx not available. Install: pip install python-docx"
            )
        
        import docx
        
        try:
            doc = docx.Document(file_path)
        except Exception as e:
            from ..exceptions import ExtractionError
            raise ExtractionError(file_path, f"Failed to open document: {e}")
        
        text_parts = []
        table_count = 0
        
        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Preserve heading styles
                if para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading ', '')
                    text = f"{'#' * int(level) if level.isdigit() else '#'} {text}"
                text_parts.append(text)
        
        # Extract tables
        for table in doc.tables:
            table_count += 1
            table_text = self._extract_table(table)
            if table_text:
                text_parts.append(table_text)
        
        # Try to get document properties for page count estimate
        page_count = self._estimate_page_count(doc)
        
        return {
            'text': '\n\n'.join(text_parts),
            'page_count': page_count,
            'method': 'python-docx',
            'metadata': {
                'library': 'python-docx',
                'paragraph_count': len(doc.paragraphs),
                'table_count': table_count
            }
        }
    
    def _extract_table(self, table) -> str:
        """Extract text from a table"""
        lines = []
        
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', ' ')
                cells.append(cell_text)
            if any(cells):
                lines.append(' | '.join(cells))
        
        return '\n'.join(lines) if lines else ''
    
    def _estimate_page_count(self, doc) -> int:
        """
        Estimate page count based on content.
        
        Word documents don't store page count directly.
        Uses rough estimate: ~3000 chars per page.
        """
        total_chars = sum(len(p.text) for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    total_chars += len(cell.text)
        
        # Rough estimate: 3000 chars per page
        return max(1, total_chars // 3000)
