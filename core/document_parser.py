"""
Document Parser Module

Parses PDF, DOCX, and TXT files for analysis.
"""

import os
from pathlib import Path
from typing import Dict, Any, Optional, List
from utils.logger_utils import get_logger

logger = get_logger("DocumentParser")


class DocumentParser:
    """Parse various document formats"""

    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']

    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a document and extract text content.

        Args:
            file_path: Path to the document

        Returns:
            Dictionary with extracted content and metadata
        """
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext not in self.supported_formats:
            return {
                'success': False,
                'error': f'Unsupported format: {ext}',
                'content': '',
                'metadata': {}
            }

        try:
            if ext == '.pdf':
                return self._parse_pdf(path)
            elif ext in ['.docx', '.doc']:
                return self._parse_docx(path)
            elif ext == '.txt':
                return self._parse_txt(path)
        except Exception as e:
            logger.error(f"Failed to parse {file_path}: {e}")
            return {
                'success': False,
                'error': str(e),
                'content': '',
                'metadata': {}
            }

    def _parse_pdf(self, path: Path) -> Dict[str, Any]:
        """Parse PDF file"""
        try:
            import PyPDF2
        except ImportError:
            # Fallback to pdfplumber
            try:
                import pdfplumber
                return self._parse_pdf_pdfplumber(path)
            except ImportError:
                return {
                    'success': False,
                    'error': 'PDF parsing requires PyPDF2 or pdfplumber. Install with: pip install PyPDF2',
                    'content': '',
                    'metadata': {}
                }

        content_parts = []
        metadata = {
            'format': 'PDF',
            'filename': path.name,
            'pages': 0
        }

        with open(path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            metadata['pages'] = len(reader.pages)

            # Extract document info
            if reader.metadata:
                metadata['title'] = reader.metadata.get('/Title', '')
                metadata['author'] = reader.metadata.get('/Author', '')

            # Extract text from each page
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    content_parts.append(f"--- Page {i+1} ---\n{text}")

        content = '\n\n'.join(content_parts)
        metadata['char_count'] = len(content)
        metadata['word_count'] = len(content.split())

        logger.info(f"Parsed PDF: {path.name}, {metadata['pages']} pages, {metadata['word_count']} words")

        return {
            'success': True,
            'content': content,
            'metadata': metadata
        }

    def _parse_pdf_pdfplumber(self, path: Path) -> Dict[str, Any]:
        """Parse PDF using pdfplumber (fallback)"""
        import pdfplumber

        content_parts = []
        metadata = {
            'format': 'PDF',
            'filename': path.name,
            'pages': 0
        }

        with pdfplumber.open(path) as pdf:
            metadata['pages'] = len(pdf.pages)

            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    content_parts.append(f"--- Page {i+1} ---\n{text}")

        content = '\n\n'.join(content_parts)
        metadata['char_count'] = len(content)
        metadata['word_count'] = len(content.split())

        return {
            'success': True,
            'content': content,
            'metadata': metadata
        }

    def _parse_docx(self, path: Path) -> Dict[str, Any]:
        """Parse DOCX file"""
        try:
            from docx import Document
        except ImportError:
            return {
                'success': False,
                'error': 'DOCX parsing requires python-docx. Install with: pip install python-docx',
                'content': '',
                'metadata': {}
            }

        doc = Document(path)
        content_parts = []
        metadata = {
            'format': 'DOCX',
            'filename': path.name,
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables)
        }

        # Extract core properties
        if doc.core_properties:
            metadata['title'] = doc.core_properties.title or ''
            metadata['author'] = doc.core_properties.author or ''

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                content_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append(' | '.join(row_text))
            if table_text:
                content_parts.append('\n'.join(table_text))

        content = '\n\n'.join(content_parts)
        metadata['char_count'] = len(content)
        metadata['word_count'] = len(content.split())

        logger.info(f"Parsed DOCX: {path.name}, {metadata['paragraphs']} paragraphs, {metadata['word_count']} words")

        return {
            'success': True,
            'content': content,
            'metadata': metadata
        }

    def _parse_txt(self, path: Path) -> Dict[str, Any]:
        """Parse TXT file"""
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252']
        content = None

        for encoding in encodings:
            try:
                content = path.read_text(encoding=encoding)
                break
            except UnicodeDecodeError:
                continue

        if content is None:
            return {
                'success': False,
                'error': 'Failed to decode text file',
                'content': '',
                'metadata': {}
            }

        metadata = {
            'format': 'TXT',
            'filename': path.name,
            'char_count': len(content),
            'word_count': len(content.split()),
            'line_count': len(content.splitlines())
        }

        logger.info(f"Parsed TXT: {path.name}, {metadata['word_count']} words")

        return {
            'success': True,
            'content': content,
            'metadata': metadata
        }

    def chunk_content(self, content: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """
        Split content into overlapping chunks for processing.

        Args:
            content: Text content to chunk
            chunk_size: Maximum characters per chunk
            overlap: Overlap between chunks

        Returns:
            List of text chunks
        """
        if len(content) <= chunk_size:
            return [content]

        chunks = []
        start = 0

        while start < len(content):
            end = start + chunk_size

            # Try to break at sentence boundary
            if end < len(content):
                # Look for sentence end
                for sep in ['. ', '.\n', '\n\n']:
                    last_sep = content[start:end].rfind(sep)
                    if last_sep > chunk_size // 2:
                        end = start + last_sep + len(sep)
                        break

            chunks.append(content[start:end].strip())
            start = end - overlap

        return chunks


# Global parser instance
_parser = None


def get_parser() -> DocumentParser:
    """Get global parser instance"""
    global _parser
    if _parser is None:
        _parser = DocumentParser()
    return _parser


def parse_document(file_path: str) -> Dict[str, Any]:
    """Convenience function to parse a document"""
    return get_parser().parse(file_path)


if __name__ == "__main__":
    import tempfile

    print("=" * 60)
    print("DOCUMENT PARSER TEST")
    print("=" * 60)

    parser = DocumentParser()

    # Create test file
    test_content = """
    Undang-Undang Republik Indonesia
    Nomor 11 Tahun 2008
    Tentang Informasi dan Transaksi Elektronik

    BAB I - KETENTUAN UMUM

    Pasal 1
    Dalam Undang-Undang ini yang dimaksud dengan:
    1. Informasi Elektronik adalah satu atau sekumpulan data elektronik.
    2. Transaksi Elektronik adalah perbuatan hukum yang dilakukan dengan menggunakan komputer.
    """

    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
        f.write(test_content)
        test_file = f.name

    print(f"\nTest file: {test_file}")

    # Parse
    result = parser.parse(test_file)

    if result['success']:
        print("\n✓ Parse successful!")
        print(f"\nMetadata:")
        for key, value in result['metadata'].items():
            print(f"  {key}: {value}")

        print(f"\nContent preview ({len(result['content'])} chars):")
        print(result['content'][:300])

        # Test chunking
        chunks = parser.chunk_content(result['content'], chunk_size=200, overlap=50)
        print(f"\nChunking: {len(chunks)} chunks")
    else:
        print(f"\n✗ Parse failed: {result['error']}")

    # Cleanup
    os.unlink(test_file)

    print("\n" + "=" * 60)
    print("TEST COMPLETE")
    print("=" * 60)
